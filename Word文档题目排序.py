from docx import Document
from pypinyin import pinyin, Style
import os

def is_chinese_char(char):
    return '\u4e00' <= char <= '\u9fff'

def get_target_char(text, start_index, config):
    index = start_index
    chars_pinyin = []
    
    if config['mode'] == 'symbol':
        # 符号定位模式
        while index < len(text):
            if text[index] in config['value']:
                index += 1
                break
            index += 1
    else:
        # 字符数跳过模式
        index += config['value']
    
    while index < len(text):
        if is_chinese_char(text[index]):
            char_pinyin = pinyin(text[index], style=Style.NORMAL)[0][0]
            if not chars_pinyin or chars_pinyin[-1][1] == char_pinyin:
                chars_pinyin.append((text[index], char_pinyin))
            else:
                break
        index += 1
    return chars_pinyin, index

def add_pinyin_to_doc(input_path, output_path, config):
    try:
        doc = Document(input_path)
        lines = [para.text for para in doc.paragraphs]

        questions = []
        current_question = []
        pinyin_char = ""

        for line in lines:
            if line.strip() and line.strip()[0].isdigit():
                if current_question:
                    questions.append((pinyin_char, current_question))
                    current_question = []
                question_text = line.strip()
                target_chars_pinyin, target_index = get_target_char(question_text, 0, config)
                if target_chars_pinyin:
                    pinyin_char = ''.join(pinyin for _, pinyin in target_chars_pinyin)
                    new_question_text = question_text[:target_index]  # 修改此处
                    for char, pinyin in target_chars_pinyin:
                        new_question_text += f"{char}"
                    new_question_text += question_text[target_index:]
                    current_question.append(new_question_text)
                else:
                    current_question.append(question_text)
            else:
                current_question.append(line.strip())
        if current_question:
            questions.append((pinyin_char, current_question))

        questions.sort(key=lambda x: x[0])

        doc = Document()

        for _, question in questions:
            for line in question:
                doc.add_paragraph(line)

        doc.save(output_path)
        print(f"文件已成功保存到: {output_path}")

    except FileNotFoundError:
        print(f"错误：找不到输入文件 '{input_path}'，请检查路径是否正确。")
    except ValueError:
        print("错误：请确保输入文件是有效的Word文档(.docx格式)")
    except Exception as e:
        print(f"发生错误: {e}")

if __name__ == "__main__":
    input_path = input("请输入文档路径(例如: C:/文档.txt): ").strip().strip('"').strip("'")
    output_path = input("请输入保存路径(例如: C:/标注后的文档.docx): ").strip().strip('"').strip("'")
    
    input_path = os.path.normpath(input_path)
    output_path = os.path.normpath(output_path)
    
    if not os.path.exists(input_path):
        print('错误：输入文件不存在！')
        exit(1)
    
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    while True:
        mode = input("请选择模式(1:符号定位 2:字符数跳过): ").strip()
        if mode not in ('1', '2'):
            print("无效模式选择，请输入1或2")
            continue
        break

    if mode == '1':
        while True:
            symbol_input = input("请输入分隔符号（多个用逗号分隔，例如：.，、）: ").strip()
            if not symbol_input:
                print("输入不能为空，请重新输入")
                continue
            symbols = [s.strip() for s in symbol_input.split(',')]
            if any(is_chinese_char(c) for s in symbols for c in s):
                print("符号不能包含中文字符，请重新输入")
                continue
            break
        add_pinyin_to_doc(input_path, output_path, {'mode': 'symbol', 'value': symbols})
    else:
        while True:
            skip_num = input("请输入要跳过的字符数: ").strip()
            if not skip_num.isdigit():
                print("请输入有效数字")
                continue
            break
        add_pinyin_to_doc(input_path, output_path, {'mode': 'skip', 'value': int(skip_num)})